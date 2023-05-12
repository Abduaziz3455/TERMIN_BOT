import docx
import re
import json

pattern = r'([A-Za-z0-9\s’\'-]+?)\s*-\s*([\s\S]+?)(?=\n\n[A-Za-z0-9\s’\'-]+?\s*-\s*|$)'
docs = ['Elektrotexnika', 'Hisoblash usullarini algoritmlash', 'Sxematexnika']


def word_atama():
    data = {}
    for i in docs:
        doc = docx.Document(f'handlers/users/{i}.docx')
        for pg in doc.paragraphs:
            matches = re.findall(pattern, pg.text)
            for match in matches:
                key = match[0].strip().lower()
                value = match[1].strip().lower()
                if len(key) > 3 and value:
                    data[key] = value
    return data

# for paragraph in doc.paragraphs:
#     # Iterate over the runs in the paragraph
#     for run in paragraph.runs:
#         # Check if the run is bold
#         if run.bold:
#             print(run.text)
